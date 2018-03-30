import docx
from docx import Document
from docx.shared import Inches
import xlwings as xw
import os
import pythoncom
import datetime
import random
import docx
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
import json
from .models import Project

def replace_text(old_text, new_text,document):
    for p in document.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in inline:
                if old_text in i.text:
                    text = i.text.replace(old_text, new_text)
                    i.text = text
    for table in document.tables:
        for column in table.columns:
            for cell in column.cells:
                for p in cell.paragraphs:
                    if old_text in p.text:
                        inline = p.runs
                        for i in inline:
                            if old_text in i.text:
                                text = i.text.replace(old_text, new_text)
                                i.text = text
                for table_sec in cell.tables:
                    for column_sec in table_sec.columns:
                        for cell_sec in column_sec.cells:
                            for p_sec in cell_sec.paragraphs:
                                if old_text in p_sec.text:
                                    inline = p_sec.runs
                                    for i in inline:
                                        if old_text in i.text:
                                            text = i.text.replace(old_text, new_text)
                                            i.text = text
def replace_word(project,document):
    replace_text("信息14", str(project.floorSpace),document)
    replace_text("信息15", str(project.managementSpace),document)
    replace_text("信息11", str(project.totalInvestment),document)
    replace_text("信息12", str(project.environmentalProtectionInvestment),document)
    replace_text("信息13", str(project.environmentalProtectionInvestment/project.totalInvestment),document)
    replace_text("信息21", str(project.investmentTime),document)
    replace_text("信息19",str( project.dayWorkTime),document)
    replace_text("信息20", str(project.yearWorkTime),document)
    replace_text("员信1", str(project.nonAccommodationNum+project.accommodationNum), document)# 员信1
    # 员信2 生活用水
    replace_text("信息42", project.domesticSewageGo, document)#信息42
    replace_text("信息44", project.pollutantHoldingWaterBody, document)#信息44
    replace_text("信息38", project.groundwaterQualityStandard, document) #信息38
    replace_text("信息36", project.groundwaterArea, document)  #信息36
    replace_text("信息39", project.groundwaterBodyNumber, document)  #信息39
    replace_text("信息46", project.surfaceWaterFunction, document)  #    信息46
    replace_text("信息43", project.domesticSewageEnvironmentImpactAnalysis, document)  #信息43
    replace_text("信息45", project.surfaceWaterQualityStandard + "水体功能",document)  #信息45 project.surfaceWaterQualityStandard + "水体功能"
    replace_text("信息38", project.groundwaterQualityStandard + "水体功能",document)  #信息38
    emissionStandard = json.loads(project.emissionStandard)
    replace_text("气排标准1", emissionStandard[0]["standard"], document)  # 信息38
    replace_text("信息36", str(project.groundwaterArea),document)
    replace_text("信息41",str(project.besideWaterTreatmentPlant),document)
    replace_text("信息28", str(project.latitude),document)
    replace_text("信息29", str(project.longtitude),document)
    replace_text("信息30", str(project.east),document)
    replace_text("信息31", str(project.south),document)
    replace_text("信息32", str(project.west),document)
    replace_text("信息33", str(project.north),document)
    replace_text("信息23", project.constructionScale,document)
    replace_text("信息35", str(project.soundEnvironmentStandard),document)
    replace_text("信息24", str(project.noiseEquipment),document)
    replace_text("信息18", str(project.accommodationNum+project.nonAccommodationNum),document)
    replace_text("信息10", project.address,document)
    replace_text("信息3", project.projectName,document)
    replace_text("信息4", project.EAcompanyName,document)
    replace_text("信息6", project.corporateName,document)
    replace_text("信息7", project.contacts, document)
    replace_text("信息8", project.telephone, document)
    replace_text("信息9", project.postalCode,document)
    replace_text("信息1 ", project.NEIType,document)
    print("finish")
def table_creater(name,wb,columns,cell,i,column,styles):
    if (cell.text == name):
        sheet = wb.sheets[name]
        rng = sheet.range('B1').expand('table')
        nrows = rng.rows.count
        rng = sheet.range('A1').expand('table')
        if(nrows < rng.rows.count):
            nrows = rng.rows.count
        newtable=cell.add_table(nrows,columns)
        for s in styles:
            if s.name == "Table Grid":
                print(s.name)
                newtable.style=s
        for a in range(0, nrows):
            for b in range(0, columns):
                print(str(a), str(b), sheet[a, b].value)
                if (sheet[a, b].value == None):
                    newtable.cell(a, b).text = " "
                else:
                    newtable.cell(a, b).text = str(sheet[a, b].value)
        cell.merge(column.cells[i])
        cell.merge(column.cells[i-2])
        print(cell.text)
        print(column.cells[i-2].text)
def create_tables(project,document):
    pythoncom.CoInitialize()
    #app = xw.App(add_book=False)
    #wb = app.books.open(project.projectName+".xlsx")
    styles = document.styles
    for table in document.tables:
        for column in table.columns:
            i = 0
            for cell in column.cells:
                i = i+1
                if(cell.text=="工程组成"):
                    otherEngineering = json.loads(project.otherEngineering)
                    environmentalEngineering = json.loads(project.environmentalEngineering)
                    rows = 5+len(environmentalEngineering)
                    columns = 3
                    newtable = cell.add_table(rows, columns)
                    for s in styles:
                        if s.name == "Table Grid":
                            print(s.name)
                            newtable.style = s
                            newtable.cell(0, 0).text = '项目'
                            newtable.cell(0, 1).text = '内容'
                            newtable.cell(0, 2).text = '用途'
                    for t in range(1, rows):
                        if(t <= 4):
                            newtable.cell(t,0).text = otherEngineering[t-1]["project"]
                            newtable.cell(t,1).text = otherEngineering[t-1]["content"]
                            newtable.cell(t,2).text = otherEngineering[t-1]["use"]
                        else:
                            newtable.cell(t, 0).text = environmentalEngineering[t-5]["project"]+str(t-4)
                            newtable.cell(t, 1).text = environmentalEngineering[t-5]["content"]
                            newtable.cell(t, 2).text = environmentalEngineering[t-5]["use"]
                    cell.merge(column.cells[i])
                    cell.merge(column.cells[i - 2])
                    print(cell.text)
                    print(column.cells[i - 2].text)
                if(cell.text == "敏感点表"):
                    sensitiveInfoWater = json.loads(project.sensitiveInfoWater)
                    sensitiveInfoReserve = json.loads(project.sensitiveInfoReserve)
                    sensitiveInfoHouse = json.loads(project.sensitiveInfoHouse)
                    rows = len(sensitiveInfoWater) + len(sensitiveInfoReserve) + len(sensitiveInfoHouse) + 3
                    columns = 5
                    newtable = cell.add_table(rows, columns)
                    for s in styles:
                        if s.name == "Table Grid":
                            print(s.name)
                            newtable.style = s
                    newtable.cell(0, 0).text = "环境要素"
                    newtable.cell(0, 1).text = "环境敏感点"
                    newtable.cell(0, 2).text = "方位"
                    newtable.cell(0, 3).text = "距离(m)"
                    newtable.cell(0, 4).text = "环境保护目标"
                    for t in range(1, rows):
                        if(t <= len(sensitiveInfoWater)):
                            newtable.cell(t,0).text = sensitiveInfoWater[t-1]["environmentalElements"]
                            newtable.cell(t,1).text = sensitiveInfoWater[t-1]["environmentalSensitivePoint"]
                            newtable.cell(t,2).text = sensitiveInfoWater[t-1]["orientation"]
                            newtable.cell(t,3).text = sensitiveInfoWater[t - 1]["distance"]
                            newtable.cell(t,4).text = sensitiveInfoWater[t - 1]["environmentalObjective"]
                        elif(t>len(sensitiveInfoWater) and t<=len(sensitiveInfoWater)+2):
                            if(t-len(sensitiveInfoWater))==1:
                                newtable.cell(t, 0).text = "大气环境"
                                newtable.cell(t, 1).text = ""
                                newtable.cell(t, 2).text = "---"
                                newtable.cell(t, 3).text = "---"
                                newtable.cell(t, 4).text = "《环境空气质量标准》（GB3095-2012）二级标准；"
                                newtable.cell(t, 0).merge(newtable.cell(t, 1))
                            if (t - len(sensitiveInfoWater)) == 2:
                                newtable.cell(t, 0).text = "声环境"
                                newtable.cell(t, 1).text = ""
                                newtable.cell(t, 2).text = "---"
                                newtable.cell(t, 3).text = "---"
                                newtable.cell(t, 4).text = "《声环境质量标准》（GB3096-2008）中的2类标准；"
                                newtable.cell(t, 0).merge(newtable.cell(t, 1))
                        elif(t > len(sensitiveInfoWater) + 2 and t <= len(sensitiveInfoWater) + 2 + len(sensitiveInfoReserve)):
                            newtable.cell(t, 0).text = sensitiveInfoReserve[t - len(sensitiveInfoWater) - 3]["environmentalElements"]
                            newtable.cell(t, 1).text = ""
                            newtable.cell(t, 2).text = sensitiveInfoReserve[t - len(sensitiveInfoWater) - 3]["orientation"]
                            newtable.cell(t, 3).text = sensitiveInfoReserve[t - len(sensitiveInfoWater) - 3]["distance"]
                            newtable.cell(t, 4).text = sensitiveInfoReserve[t - len(sensitiveInfoWater) - 3]["environmentalObjective"]
                            newtable.cell(t, 0).merge(newtable.cell(t, 1))
                        elif (t > len(sensitiveInfoWater) + 2 + len(sensitiveInfoReserve)):
                            newtable.cell(t, 0).text = sensitiveInfoHouse[t - len(sensitiveInfoWater) - 3 - len(sensitiveInfoReserve)]["environmentalElements"]
                            newtable.cell(t, 1).text = ""
                            newtable.cell(t, 2).text = sensitiveInfoHouse[t - len(sensitiveInfoWater) - 3 - len(sensitiveInfoReserve)]["orientation"]
                            newtable.cell(t, 3).text = sensitiveInfoHouse[t - len(sensitiveInfoWater) - 3 - len(sensitiveInfoReserve)]["distance"]
                            newtable.cell(t, 4).text = sensitiveInfoHouse[t - len(sensitiveInfoWater) - 3 - len(sensitiveInfoReserve)]["environmentalObjective"]
                            newtable.cell(t, 0).merge(newtable.cell(t, 1))
                    cell.merge(column.cells[i])
                    cell.merge(column.cells[i - 2])
                # wb.close()
                # app.quit()
                if (cell.text == "废气排放标准表"):
                    emissionStandard = json.loads(project.emissionStandard)
                    rows = len(emissionStandard) + 1
                    columns = 5
                    newtable = cell.add_table(rows, columns)
                    for s in styles:
                        if s.name == "Table Grid":
                            print(s.name)
                            newtable.style = s
                    newtable.cell(0, 0).text = "污染物"
                    newtable.cell(0, 1).text = "最高允许排放浓度（mg / m3）"
                    newtable.cell(0, 2).text = "最高允许排放速率（kg / h） "
                    newtable.cell(0, 3).text = "无组织排放监控点浓度限值（mg / m3）"
                    newtable.cell(0, 4).text = "标准"
                    for t in range(1, rows):
                        newtable.cell(t,0).text = emissionStandard[t-1]["pollutant"]
                        newtable.cell(t,1).text = emissionStandard[t-1]["maximumAllowableEmissionConcentration"]
                        newtable.cell(t,2).text = emissionStandard[t-1]["maximumAllowableEmissionRate"]
                        newtable.cell(t,3).text = emissionStandard[t - 1]["emissionMonitoring"]
                        newtable.cell(t,4).text = emissionStandard[t - 1]["standard"]
                    cell.merge(column.cells[i])
                    cell.merge(column.cells[i - 2])



def createWord(request,projectName):
    document = Document('评估报告(新建模板).docx')
    project = Project.objects.get(projectName=projectName)
    print(projectName)
    #replace_word(project,document)
    create_tables(project,document)
    document.save('demo.docx')
